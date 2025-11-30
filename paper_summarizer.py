import sys
import ollama
from docx import Document
from utils import extract_text_from_pdf

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python paper_summarizer.py <pdf_path>")
        sys.exit(1)

    pdf_path = sys.argv[1]
    text = extract_text_from_pdf(pdf_path)

    if not text.strip():
        print("No text extracted from PDF.")
        sys.exit(1)

    # Summarize with Ollama
    prompt = f"Summarize the following academic paper in 200-300 words, focusing on key findings, methods, and conclusions:\n\n{text[:4000]}"  # Limit text to avoid token limits
    response = ollama.generate(model='llama3:8b', prompt=prompt)
    summary = response['response']

    # Create .docx
    doc = Document()
    doc.add_heading('Paper Summary', 0)
    doc.add_paragraph(summary)
    output_path = f"summaries/{pdf_path.split('/')[-1].replace('.pdf', '_summary.docx')}"
    doc.save(output_path)

    print(f"Summary saved to {output_path}")
