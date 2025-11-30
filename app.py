"""Give It A Summary - Main Application Module."""

from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import FileResponse
import ollama
from docx import Document
from backend.app.utilities.extraction import extract_text
import os

app = FastAPI(title="Give-It-A-Summary API", description="Upload a PDF or TXT paper and get a summarized DOCX back.")

@app.post("/summarize")
async def summarize_paper(file: UploadFile = File(...)):
    if not (file.filename.endswith('.pdf') or file.filename.endswith('.txt') or file.filename.endswith(('.xlsx', '.xls'))):
        raise HTTPException(status_code=400, detail="Only PDF, TXT, XLSX, or XLS files are allowed.")

    if file.filename.endswith('.pdf'):
        # Save temp PDF
        temp_path = f"/tmp/{file.filename}"
        with open(temp_path, "wb") as f:
            content = await file.read()
            f.write(content)
        text = extract_text_from_pdf(temp_path)
        os.remove(temp_path)
    elif file.filename.endswith('.txt'):
        content = await file.read()
        text = content.decode('utf-8')
    elif file.filename.endswith(('.xlsx', '.xls')):
        # Save temp Excel
        temp_path = f"/tmp/{file.filename}"
        with open(temp_path, "wb") as f:
            content = await file.read()
            f.write(content)
        text = extract_text_from_excel(temp_path)
        os.remove(temp_path)

    if not text.strip():
        raise HTTPException(status_code=400, detail="No text extracted from the file.")

    # Summarize with Ollama
    prompt = f"Summarize the following academic paper in 200-300 words, focusing on key findings, methods, and conclusions:\n\n{text[:4000]}"
    response = ollama.generate(model='llama3:8b', prompt=prompt)
    summary = response['response']

    # Create DOCX
    doc = Document()
    doc.add_heading('Paper Summary', 0)
    doc.add_paragraph(summary)
    output_filename = f"{file.filename.replace('.pdf', '').replace('.txt', '').replace('.xlsx', '').replace('.xls', '')}_summary.docx"
    output_path = f"/tmp/{output_filename}"
    doc.save(output_path)

    return FileResponse(output_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=output_filename)