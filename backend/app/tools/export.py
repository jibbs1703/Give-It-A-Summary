"""Give It A Summary - Document Export Tools Module."""

import os
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from pydantic import BaseModel, Field

from app.utilities.logs import get_logger

logger = get_logger(__name__)


class DocumentExportInputs(BaseModel):
    """Pydantic model for document export inputs."""

    summary: str
    original_filename: str
    word_count: int
    style: str
    output_path: str | None = None


class DocumentExportResult(BaseModel):
    """Result of document export operation."""

    file_path: str = Field(..., description="Path to the generated document")
    file_size: int = Field(..., description="Size of the generated document in bytes")
    success: bool = Field(default=True, description="Whether export was successful")
    error_message: str | None = Field(None, description="Error message if export failed")


def setup_document_styles(doc: Document) -> None:
    """Set up document styles."""
    doc.styles.add_style("CustomTitle", WD_STYLE_TYPE.PARAGRAPH).font.size = Pt(24)
    doc.styles["CustomTitle"].font.bold = True
    doc.styles["CustomTitle"].font.name = "Arial"

    doc.styles.add_style("CustomSubtitle", WD_STYLE_TYPE.PARAGRAPH).font.size = Pt(14)
    doc.styles["CustomSubtitle"].font.italic = True
    doc.styles["CustomSubtitle"].font.name = "Arial"

    header_style = doc.styles.add_style("CustomHeader", WD_STYLE_TYPE.PARAGRAPH)
    header_style.font.size = Pt(16)
    header_style.font.bold = True
    header_style.font.name = "Arial"
    header_style.paragraph_format.space_after = Pt(12)

    body_style = doc.styles.add_style("CustomBody", WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.size = Pt(11)
    body_style.font.name = "Times New Roman"
    body_style.paragraph_format.space_after = Pt(8)
    body_style.paragraph_format.line_spacing = 1.15

    meta_style = doc.styles.add_style("CustomMetadata", WD_STYLE_TYPE.PARAGRAPH)
    meta_style.font.size = Pt(10)
    meta_style.font.name = "Arial"
    meta_style.paragraph_format.space_after = Pt(6)


def export_to_word_document(inputs: DocumentExportInputs) -> DocumentExportResult:
    """
    Export summary to a clean Word document.

    Args:
        inputs: Document export parameters

    Returns:
        DocumentExportResult with file path and metadata
    """
    try:
        if inputs.output_path:
            output_path = Path(inputs.output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
        else:
            temp_fd, temp_path = tempfile.mkstemp(suffix=".docx", prefix="summary_")
            os.close(temp_fd)
            output_path = Path(temp_path)

        logger.info("Creating Word document at: %s", output_path)

        doc = Document()
        setup_document_styles(doc)

        header = doc.add_paragraph(style="CustomHeader")
        header.add_run("Summary")

        if inputs.style == "bullets":
            for line in inputs.summary.strip().splitlines():
                stripped = line.strip().lstrip("-*• ").strip()
                if stripped:
                    p = doc.add_paragraph(style="CustomBody")
                    p.add_run(f"\u2022 {stripped}")
        else:
            for para in inputs.summary.strip().split("\n\n"):
                para = para.strip()
                if para:
                    p = doc.add_paragraph(style="CustomBody")
                    p.add_run(para.replace("\n", " "))

        doc.save(str(output_path))

        file_size = output_path.stat().st_size

        logger.info("Word document created successfully: %s (%d bytes)", output_path, file_size)

        return DocumentExportResult(file_path=str(output_path), file_size=file_size, success=True)

    except Exception as e:  # noqa: BLE001
        error_msg = f"Document export failed: {str(e)}"
        logger.error("Document export error: %s (%s)", str(e), type(e).__name__)

        if "output_path" in locals() and output_path.exists():
            try:
                output_path.unlink()
            except Exception as cleanup_error:  # noqa: BLE001
                logger.warning(
                    "Failed to cleanup failed export file %s: %s", output_path, str(cleanup_error)
                )

        return DocumentExportResult(
            file_path="", file_size=0, success=False, error_message=error_msg
        )


def cleanup_temp_file(file_path: str) -> bool:
    """
    Clean up a temporary document file.

    Args:
        file_path: Path to the file to clean up

    Returns:
        True if cleanup was successful
    """
    try:
        path = Path(file_path)
        if path.exists() and path.is_file():
            path.unlink()
            logger.info("Cleaned up temporary document file: %s", file_path)
            return True
    except Exception as e:  # noqa: BLE001
        logger.warning("Failed to cleanup temporary file %s: %s", file_path, str(e))

    return False
